from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook

#替换word中的关键词
def info_update(doc, old_info, new_info):
    #遍历替换所有段落中的关键词
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info)
    #遍历替换所有表格中的关键词
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info)
#读取excel中的信息
wb = load_workbook('简历信息.xlsx')
ws = wb.active
#根据简历信息中的每行生成一个简历，即为每个公司生成一份专属简历
for row in range(2, ws.max_row + 1):
    doc = Document('简历模板.docx')
    for col in range(1, ws.max_column + 1):
        #简历信息中的标题，对应word中的关键词
        old_info = str(ws.cell(row=1, column=col).value)
        #简历信息中的替换的具体公司关键词
        new_info = str(ws.cell(row=row, column=col).value)
        #执行替换
        info_update(doc, old_info, new_info)
    #插入电子签名
    #在word中的第2个表格的（0，1）单元格插入图片
    table = doc.tables[1]
    run = table.cell(0,1).paragraphs[0].add_run('')
    run.add_picture('qianzi-binary.png', width=Inches(1), height=Inches(0.5))
    #基于更新后的简历模板生成以对应公司命名的简历
    com_name = str(ws.cell(row=row, column=1).value)
    doc.save(f'{com_name}简历.docx')
